from docx import Document
from docx.shared import Pt, Inches
import pdfkit
import os
from typing import Dict
import re
import subprocess

class CourseExporter:
    def __init__(self):
        """Initialize the course exporter."""
        # Configure pdfkit with wkhtmltopdf path
        self.wkhtmltopdf_path = self._find_wkhtmltopdf()
        if self.wkhtmltopdf_path:
            self.config = pdfkit.configuration(wkhtmltopdf=self.wkhtmltopdf_path)
        else:
            self.config = None

    def _find_wkhtmltopdf(self) -> str:
        """Find wkhtmltopdf executable path."""
        # Common installation paths
        possible_paths = [
            r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe',
            r'C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe',
            r'C:\wkhtmltopdf\bin\wkhtmltopdf.exe'
        ]
        
        # Check if wkhtmltopdf is in PATH
        try:
            result = subprocess.run(['where', 'wkhtmltopdf'], 
                                 capture_output=True, 
                                 text=True, 
                                 check=True)
            if result.stdout:
                return result.stdout.strip()
        except:
            pass
        
        # Check common installation paths
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None

    def _sanitize_filename(self, filename: str) -> str:
        """Remove invalid characters from filename."""
        # Remove quotes and other invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        # Replace multiple spaces with single space
        filename = re.sub(r'\s+', ' ', filename)
        return filename.strip()

    def export_to_docx(self, course_content: Dict, output_path: str):
        """Export course content to DOCX format."""
        doc = Document()
        
        # Add title
        doc.add_heading(course_content["title"], 0)
        
        # Add objectives
        doc.add_heading("Learning Objectives", level=1)
        for objective in course_content["objectives"]:
            doc.add_paragraph(objective, style='List Bullet')
        
        # Add sections
        for section in course_content["sections"]:
            doc.add_heading(section["title"], level=1)
            
            # Add content
            doc.add_paragraph(section["content"])
            
            # Add summary
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(section["summary"])
            
            # Add quiz
            doc.add_heading("Quiz", level=2)
            for i, question in enumerate(section["quiz"], 1):
                doc.add_paragraph(f"Question {i}: {question['question']}")
                for option in question["options"]:
                    doc.add_paragraph(option, style='List Bullet')
                doc.add_paragraph(f"Correct Answer: {question['correct_answer']}")
        
        # Save document
        doc.save(output_path)

    def export_to_pdf(self, course_content: Dict, output_path: str):
        """Export course content to PDF format."""
        if not self.wkhtmltopdf_path:
            raise RuntimeError(
                "wkhtmltopdf not found. Please install it from https://wkhtmltopdf.org/downloads.html"
            )

        # First create a temporary HTML file
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #34495e; }}
                .objective {{ margin: 10px 0; }}
                .section {{ margin: 20px 0; }}
                .quiz {{ margin: 15px 0; }}
            </style>
        </head>
        <body>
            <h1>{course_content["title"]}</h1>
            
            <h2>Learning Objectives</h2>
            <ul>
                {"".join(f'<li class="objective">{obj}</li>' for obj in course_content["objectives"])}
            </ul>
        """
        
        for section in course_content["sections"]:
            html_content += f"""
            <div class="section">
                <h2>{section["title"]}</h2>
                <p>{section["content"]}</p>
                
                <h3>Summary</h3>
                <p>{section["summary"]}</p>
                
                <h3>Quiz</h3>
            """
            
            for i, question in enumerate(section["quiz"]):
                html_content += f"""
                <div class="quiz">
                    <p><strong>Question {i+1}:</strong> {question["question"]}</p>
                    <ul>
                        {"".join(f'<li>{opt}</li>' for opt in question["options"])}
                    </ul>
                    <p><em>Correct Answer: {question["correct_answer"]}</em></p>
                </div>
                """
            
            html_content += """
            </div>
            """
        
        html_content += """
        </body>
        </html>
        """
        
        # Write HTML to temporary file
        temp_html = "temp_course.html"
        with open(temp_html, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        try:
            # Convert HTML to PDF
            pdfkit.from_file(temp_html, output_path, configuration=self.config)
        finally:
            # Clean up temporary file
            if os.path.exists(temp_html):
                os.remove(temp_html) 